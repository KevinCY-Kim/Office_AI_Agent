# -*- coding: utf-8 -*-
"""
RND_Plan_Pipeline_v3_Final.py
"""

import os, re, json, glob, gc, torch
import numpy as np
from tqdm.auto import tqdm
from transformers import (
    AutoTokenizer, AutoModelForCausalLM, pipeline,
    AutoModelForSequenceClassification
)
from sentence_transformers import SentenceTransformer, util
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
from PyPDF2 import PdfReader

# ==============================
# 1. 경로 및 모델 설정
# ==============================
GUIDELINE_FILE = "/home/alpaco/autosry/rnd_guideline.json"
RAG_JSON_FILES = ["/home/alpaco/autosry/rag_chunks500_50.json"]
LAW_DIR = "/home/alpaco/autosry/reference_file"

E5_NAME = "intfloat/multilingual-e5-large"
GEN_NAME = "skt/A.X-4.0-Light"
NLI_MODEL = "MoritzLaurer/multilingual-MiniLMv2-L6-mnli-xnli"

DEVICE = "cuda:1" if torch.cuda.is_available() else "cpu"

# GPU 메모리 감지 → 안전 토큰 설정
GEN_MAX_NEW_TOKENS = 3000
if torch.cuda.is_available():
    free, total = torch.cuda.mem_get_info()
    if total < 12 * 1024**3:  # 12GB 미만 GPU는 토큰 수 축소
        GEN_MAX_NEW_TOKENS = 1500
torch.cuda.empty_cache()
GEN_DO_SAMPLE = False

# ==============================
# 2. GPU/GC 관리 헬퍼 함수
# ==============================
def _cuda_gc():
    """GPU/CPU 메모리 동기화 및 캐시 정리"""
    gc.collect()
    if torch.cuda.is_available():
        try:
            dev_index = int(str(DEVICE).split(":")[1]) if ":" in str(DEVICE) else 0
        except:
            dev_index = 0
        with torch.cuda.device(dev_index):
            torch.cuda.empty_cache()
            torch.cuda.ipc_collect()

# ==============================
# 3. 유틸리티 함수
# ==============================
def first_n_lines(text: str, n_chars=400):
    t = " ".join(str(text).split())
    return t[:n_chars]

def has_number(text: str):
    return bool(re.search(r"\d", text))

def clean_generated_text(text: str) -> str:
    """프롬프트 헤더, 지시문, 특수문자 제거 후 정제"""
    text = re.sub(r'#=+.*?(자동\s*문장\s*생성).*?#=+', '', text, flags=re.DOTALL)
    text = re.sub(r'^###.*$', '', text, flags=re.MULTILINE)
    text = re.sub(r'[•●▪▶◇◆□▪️▫️–]', ' ', text)
    text = re.sub(r'^\s*[-#*]+\s*', '', text, flags=re.MULTILINE)
    text = re.sub(r'\n{3,}', '\n\n', text)
    text = re.sub(r'\s+([\.,;:])', r'\1', text)
    text = re.sub(r'\s{2,}', ' ', text)
    text = re.sub(r'출력은\s*본문\s*서술문\s*형태로\s*작성하시오[.\s]*', '', text)
    return text.strip()

# ==============================
# 4. 섹션 정의
# ==============================
sections = [
    {"section": "연구기획과제의 개요",
        "role": "제안서 총괄 에디터",
        "query": "과제의 목적, 필요성, 기대효과를 논리적으로 연결하여 핵심 내용을 500자 내외로 명료하게 요약함. 평가자가 과제의 전체 구조를 한눈에 이해할 수 있도록 기술함.",
        "constraints": ["500자 내외", "논리적 연결 구조 유지", "핵심 요약 중심"],
        "output_format": "서술문"},
    {"section": "연구개발과제의 배경",
        "role": "R&D 기획 전문가",
        "query": "제공된 선행연구, 시장 동향, 정책자료를 근거로 본 과제가 추진되어야 하는 당위성과 RFP(공고문) 부합성을 논리적으로 제시함.",
        "constraints": ["데이터 근거 포함", "RFP 문항 부합성 명시"],
        "output_format": "서술문"},
    {"section": "연구개발과제의 필요성",
        "role": "산업분석가",
        "query": "데이터와 사례를 근거로 현재 기술적·산업적 문제점을 제시하고, 본 과제가 이를 해결해야 하는 필요성을 인과적으로 서술함.",
        "constraints": ["인과관계 구조", "데이터 근거 제시"],
        "output_format": "서술문"},
    {"section": "기술개발 핵심어(키워드)",
        "role": "기술 네이밍 전략가",
        "query": "과제의 정체성과 핵심 기술을 대표하는 키워드 5개를 국문·영문 공식 명칭으로 제시함. 각 용어는 국제 표준 또는 학술 정의를 근거로 하며, 선정 사유를 간단히 명시함.",
        "constraints": ["국문·영문 병기", "국제 표준 기반 정의 포함", "5개 이내"],
        "output_format": "표(키워드/영문명/정의/출처)"},
    {"section": "연구개발 목표",
        "role": "R&D PMO",
        "query": "과제의 최종 목표를 500자 내외로 명확히 기술함. 핵심 성능지표(KPI), 달성 기준(수치·단위·마일스톤), 검증 방법을 포함하며, 모호한 표현은 사용하지 않음.",
        "constraints": ["정량화된 수치 포함", "KPI 및 검증방법 명시"],
        "output_format": "서술문 + 표(KPI/단위/기준/검증방법)"},
        {"section": "연구개발 내용",
            "role": "기술총괄자(Tech Lead)",
            "query": "전체 연구 범위를 1,000자 내외로 체계적으로 기술함. 핵심 기술요소, 세부 과제 구조, 데이터 및 시스템 흐름, 성능 평가 계획을 명시함.",
            "constraints": ["1,000자 내외", "기술요소 및 평가계획 포함"],
            "output_format": "서술문 + 도식(기술흐름)"},
        {"section": "연차별 개발목표",
            "role": "PMO 리더",
            "query": "최종 목표 달성을 위한 연차별 및 기관별 개발 목표를 정량적으로 제시함. 각 연차별 KPI, 마일스톤, 검증 방법을 명확히 포함함.",
            "constraints": ["연차별 구분", "정량적 지표 포함"],
            "output_format": "표(연차/KPI/마일스톤/검증방법)"},
        {"section": "연차별 개발내용 및 범위",
        "role": "공동연구 컨소시엄 코디네이터",
        "query": "참여 기관별 역할, 책임, 연차별 산출물을 중복 및 누락 없이 기술함. 공동기관이 없는 경우 해당 항목은 생략함.",
        "constraints": ["기관별 역할 명확화", "중복/누락 금지"],
        "output_format": "표(기관/역할/산출물/책임)"},
        {"section": "추진방법 및 전략",
            "role": "총괄 아키텍트(Chief Architect)",
            "query": "핵심 기술개발 방법론, 예측 가능한 리스크와 대응 방안, 성능 검증 계획을 논리적으로 제시함. 기술의 우수성과 실현 가능성을 입증함.",
            "constraints": ["핵심 방법론 포함", "리스크 및 검증계획 명시"],
            "output_format": "서술문 + 표(리스크/대응방안)"},
        {"section": "과제 성과의 활용방안",
            "role": "사업개발 총괄(BD Head)",
            "query": "연구 성과가 실제 산업 및 시장에서 어떻게 활용될 수 있는지를 구체적으로 제시함. 목표 시장, 주요 수요처, 핵심 적용 시나리오, 기술의 차별화된 가치를 중심으로 사업화 방향을 설명함.",
            "constraints": ["시장 시나리오 포함", "Value Proposition 중심"],
            "output_format": "서술문 + 표(시장/수요처/적용시나리오)"},
        {"section": "신규사업 신설의 기대효과",
            "role": "거시경제 분석가(Macro-Economic Analyst)",
            "query": "본 과제가 국가 경제에 미치는 파급효과를 정량적 지표로 제시함. 시장 창출, 수입 대체, 수출 증대, 일자리 창출 등 거시적 효과를 수치로 증명함.",
            "constraints": ["정량적 수치 기반", "경제효과 명시"],
            "output_format": "표(지표/예상값/근거자료)"},
        {"section": "사회적 가치 창출 계획",
            "role": "사회적 가치 전략가(Social Value Strategist)",
            "query": "과제의 사회적 비전과 목표를 정의하고, 이를 달성하기 위한 구체적인 실행 로드맵을 수립함. 보건, 환경, 안전 등 사회적 가치 범주와 연계함.",
            "constraints": ["사회적 가치 범주 명시", "로드맵 포함"],
            "output_format": "서술문 + 표(목표/실행단계/성과지표)"},
        {"section": "사회적 가치창출의 기대효과",
            "role": "임팩트 평가 전문가(Impact Assessor)",
            "query": "사회적 가치 창출 계획이 실행되었을 때 예상되는 긍정적 변화를 정량적 및 정성적 임팩트 지표로 제시함. 사회적 파급효과를 객관적으로 설명함.",
            "constraints": ["정량/정성 지표 병기", "사회적 파급효과 포함"],
            "output_format": "표(지표유형/성과/측정방법)"},
        {"section": "경제적 성과창출의 기대효과",
            "role": "최고재무책임자(CFO)",
            "query": "기업 관점에서 본 과제의 재무적 성과를 구체적 수치와 함께 제시함. 예상 매출, 이익, 투자수익률(ROI), 순현재가치(NPV) 등 주요 지표를 근거와 함께 명료하게 기술함.",
            "constraints": ["재무지표 포함", "산출근거 명시"],
            "output_format": "표(지표/예상값/근거)"},
        {"section": "신규 인력 채용 계획 및 활용 방안",
            "role": "전략적 인사 파트너(Strategic HR Partner)",
            "query": "과제 수행에 필요한 핵심 인력의 채용, 배치, 교육 계획을 타임라인과 함께 제시함. 인력 확보 및 역량 극대화 방안을 구체적으로 기술함.",
            "constraints": ["타임라인 포함", "역량 강화 계획 명시"],
            "output_format": "표(직무/채용시점/교육계획)"},
        {"section": "보안등급의 분류 및 해당 사유",
            "role": "보안관리 책임자(Security Manager)",
            "query": "관련 법령 및 보안관리요령을 근거로 본 과제의 보안등급을 분류하고, 그 결정 사유를 간결하고 명확하게 기술함.",
            "constraints": ["법령 근거 포함", "사유 명시"],
            "output_format": "서술문"}
        ]

# ==============================
# 5. 법령·지침 로드 및 검색
# ==============================
CHUNK_MAX = 500
CHUNK_OVERLAP = 50
TOPK = 5
REF_WEIGHTS = {
    "행정업무의 운영 및 혁신에 관한 규정(대통령령)": 0.30,
    "행정업무의 운영 및 혁신에 관한 규정 시행규칙": 0.30,
    "국가연구개발사업 연구개발계획서": 0.20,
    "전략계획서 작성안내서": 0.10,
    "Vertical": 0.10
}
DEFAULT_WEIGHT = 0.05

def guess_weight(filename: str) -> float:
    for k, w in REF_WEIGHTS.items():
        if k.lower() in filename.lower():
            return w
    return DEFAULT_WEIGHT

def chunk_text(txt: str, max_chars=CHUNK_MAX, overlap=CHUNK_OVERLAP):
    txt = " ".join(str(txt).split())
    chunks = []
    i = 0
    while i < len(txt):
        j = min(len(txt), i + max_chars)
        chunks.append(txt[i:j])
        if j == len(txt): break
        i = max(0, j - overlap)
    return chunks

def load_reference_chunks(law_dir: str):
    _cuda_gc()
    items = []
    for p in glob.glob(os.path.join(law_dir, "*")):
        text = ""
        try:
            ext = p.split(".")[-1].lower()
            if ext == "pdf":
                reader = PdfReader(p)
                for pg in reader.pages:
                    text += (pg.extract_text() or "") + "\n"
            elif ext in ("docx", "doc"):
                d = Document(p)
                text = "\n".join([x.text for x in d.paragraphs])
            elif ext in ("rtf", "txt"):
                with open(p, "r", encoding="utf-8", errors="ignore") as f:
                    text = f.read()
        except Exception:
            continue
        base = os.path.basename(p)
        weight = guess_weight(base)
        for idx, ck in enumerate(chunk_text(text)):
            items.append({"source": base, "chunk_id": idx, "weight": weight, "text": ck})
    return items

embed_model = SentenceTransformer(E5_NAME, device=DEVICE)
REF_ITEMS = load_reference_chunks(LAW_DIR)
REF_EMBS = embed_model.encode([it["text"] for it in REF_ITEMS],
                              convert_to_tensor=True,
                              normalize_embeddings=True)
print(f"[INFO] Loaded {len(REF_ITEMS)} reference chunks.")

def search_reference(query: str, topk: int = TOPK):
    global embed_model
    q = embed_model.encode(query, convert_to_tensor=True, normalize_embeddings=True)
    sims = util.cos_sim(q, REF_EMBS)[0]
    scores = [(float(s) * (1.0 + REF_ITEMS[i]["weight"]), i) for i, s in enumerate(sims)]
    scores.sort(key=lambda x: x[0], reverse=True)
    picks = []
    for sc, idx in scores[:topk]:
        it = REF_ITEMS[idx]
        picks.append({
            "source": it["source"], "chunk_id": it["chunk_id"],
            "weight": it["weight"], "score": round(sc, 4),
            "snippet": first_n_lines(it["text"], 400)
        })
    _cuda_gc()
    return picks

# ==============================
# 6. 생성모델 로드
# ==============================
gen_tok = AutoTokenizer.from_pretrained(GEN_NAME, use_fast=False, trust_remote_code=True)
if gen_tok.pad_token_id is None:
    gen_tok.pad_token = gen_tok.eos_token

gen_model = AutoModelForCausalLM.from_pretrained(GEN_NAME, trust_remote_code=True, device_map="auto").eval()
gen_pipe = pipeline("text-generation", model=gen_model, tokenizer=gen_tok)
del gen_model
_cuda_gc()

# ==============================
# 7. 검증모델 (NLI)
# ==============================
nli_tok = AutoTokenizer.from_pretrained(NLI_MODEL)
nli_model = AutoModelForSequenceClassification.from_pretrained(NLI_MODEL).to(DEVICE).eval()
_cuda_gc()

def nli_entail_vs_contra(premise, hypothesis):
    _cuda_gc()
    inputs = nli_tok(premise, hypothesis, return_tensors="pt", truncation=True, padding=True).to(DEVICE)
    with torch.no_grad():
        logits = nli_model(**inputs).logits
    probs = torch.softmax(logits, dim=-1).squeeze(0).tolist()
    entail, neutral, contra = probs[0], probs[1], probs[2]
    return entail, contra

# ==============================
# 8. 검증 함수
# ==============================
def validate_output(section_obj, generated_text: str, refs):
    _cuda_gc()
    report = []
    cleaned = generated_text.strip()
    maxlen = 10**9
    for c in section_obj.get("constraints", []):
        m = re.search(r"(\d+)\s*자", c)
        if m:
            maxlen = int(m.group(1)) + int(0.3*int(m.group(1)))
            break
    if len(cleaned) > maxlen:
        report.append(f" 길이 초과: {len(cleaned)}자 > 허용 {maxlen}자")
    if any(k in section_obj["section"] for k in ["목표", "KPI"]) or "정량" in " ".join(section_obj.get("constraints", [])):
        if not has_number(cleaned):
            report.append("정량 지표(숫자) 미포함")

    entail_sum, contra_sum = 0.0, 0.0
    for r in refs[:3]:
        e, c = nli_entail_vs_contra(r["snippet"], cleaned)
        entail_sum += e; contra_sum += c
    score = entail_sum / (entail_sum + contra_sum + 1e-6)
    report.append(f" NLI 정합도(entailment vs contradiction): {score:.2f}")
    if not report:
        report.append(" PASS-validation")
    return report

# ==============================
# 9. 프롬프트 생성
# ==============================
def build_prompt(section_obj, project_name, depart_name, project_no, period, budget):
    _cuda_gc()
    refs = search_reference(section_obj["query"], topk=TOPK)
    ref_texts = "\n\n".join([r["snippet"] for r in refs])
    ref_texts = re.sub(r'붙임|끝|제\d+조|\(\d{8}\)', '', ref_texts)
    return f"""
# #=========== 자동 문장 생성
# 역할: {section_obj['role']}
# 작성 항목: [{section_obj['section']}]
# 세부사업명: {depart_name}
# 연구개발 과제번호: {project_no}
# 연구개발과제명: {project_name}
# 전체 연구개발기간: {period}
# 예산: {budget} 천원
# 작성 조건:
# AI 총괄 지침 (MASTER PROMPT)
#     - 당신은 지금부터 R&D 보고서의 특정 섹션을 책임지는 '전문가' 역할을 수행합니다. 당신의 유일한 임무는 주어진 역할에 완벽히 몰입하여, 아래에 제공될 **[참조 데이터]만을 근거**로 담당 섹션을 작성하는 것입니다. 다음 세 가지 대원칙을 절대적으로 준수해야 합니다.
---
#     - 제1원칙: 역할 분담 및 중복 방지의 원칙 (Principle of Role Division & Non-Redundancy)
      - **역할 몰입:** 각 섹션은 서로 다른 전문가가 작성합니다. 당신의 역할에 명시된 임무와 결과물에만 집중하고, 다른 섹션의 영역을 침범하지 마십시오.
      - **중복 금지:** 예를 들어, 재무 담당자(CFO)는 사업화 모델을 상세히 논하지 않고, 오직 재무적 성과(매출, ROI 등)에만 집중해야 합니다. 이는 보고서 전체의 논리적 명확성을 위함입니다.
---
#     - 제2원칙: 근거 기반 생성 및 환각 방지의 원칙 (Principle of Fact-Based Generation & Hallucination Prevention)
      - **유일한 진실의 원천 (Single Source of Truth):** 당신에게 제공될 **[참조 데이터]**가 유일한 정보 소스입니다. 당신이 사전에 학습한 지식이나 외부 정보를 활용하여 내용을 추론하거나 확장하는 것은 **엄격히 금지**합니다.
      - **정보 부재 시 행동 강령:** 만약 특정 내용을 작성하는 데 필요한 정보가 [참조 데이터]에 없다면, 절대 내용을 임의로 생성하지 마십시오. 대신, **"[작성을 위해 OOO에 대한 구체적인 정보가 필요합니다.]"** 와 같이 명확하게 정보의 부재를 알리십시오.
---
#     - 제3원칙: 논증적 보고서 문체 및 구조의 원칙 (Principle of Argumentative Report Style & Structure)
      - **두괄식 구성 (Claim-First Structure):** 모든 문단은 **핵심 주장이나 결론을 첫 문장에 제시**하는 두괄식으로 작성해야 합니다. 그 후에 해당 주장을 뒷받침하는 근거와 데이터를 논리적으로 제시하십시오.
      - **객관적이고 간결한 톤 (Objective & Concise Tone):** 주관적, 감정적 표현을 **엄격히 금지**하고, 오직 [참조 데이터]에 기반한 객관적인 사실과 분석만을 간결하고 명료한 문장으로 전달하십시오.
      - **인과관계 명시 (Explicit Causality):** 단순 사실 나열을 넘어, **'A이기 때문에 B가 필요하다'** 와 같이 원인과 결과, 문제와 해결책 사이의 논리적 연결고리를 명확하게 서술하여 주장의 설득력을 높여야 합니다.
---
# 세부 작성 지침:
#     - 제시된 {GUIDELINE_FILE} 가이드라인을 엄격히 준수하여 작성함.
#     - 문장은 ~함, ~음, 명사로 마무리함(문장 종결 통일).
#     - 기술적 연관성이 낮은 미사여구 배제.
#     - 구체적인 규격/범위 포함.
#     - 아래 근거를 반영함:
# {ref_texts}
#     - 반드시 {RAG_JSON_FILES}의 작성 방식과 구성을 참고하여 작성함.
#     - 문단마다 핵심 키워드 포함, 문장 길이/시작 다양화, 중복 표현 회피.
#     - 전문적이면서 친화적인 톤.
#     - 모든 전문용어/약어에는 주석(full name) 표기.
# 요청된 형식(힌트): {section_obj['query']}
#=========== 출력
""".strip()

# ==============================
# 10. 생성 실행
# ==============================
def generate_text_batch(prompts, batch_size=2):
    _cuda_gc()
    outputs = []
    for i in tqdm(range(0, len(prompts), batch_size), desc="Generate"):
        chunk = prompts[i:i + batch_size]
        out = gen_pipe(chunk, batch_size=batch_size, max_new_tokens=GEN_MAX_NEW_TOKENS, do_sample=GEN_DO_SAMPLE)
        for o in out:
            txt = o[0]["generated_text"] if isinstance(o, list) else o["generated_text"]
            outputs.append(txt)
    return outputs

# ==============================
# 11. DOCX 렌더링
# ==============================
def add_reference_section(doc, all_refs):
    doc.add_page_break()
    doc.add_heading("근거 법령 및 참고 문서 목록", level=1)
    seen = {}
    for r in all_refs:
        name = r['source']
        seen[name] = seen.get(name, 0) + 1
    for name, cnt in seen.items():
        doc.add_paragraph(f"- {name} (참조 {cnt}회)")

def render_doc(project_name, depart_name, project_no, period, budget):
    _cuda_gc()
    doc = Document()
    style = doc.styles['Normal']
    font = style.font
    font.name = '맑은 고딕'
    font.size = Pt(11)
    style._element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')

    doc.add_heading("R&D 국가연구개발계획서", 0)
    doc.add_paragraph(f"세부사업명: {depart_name}")
    doc.add_paragraph(f"과제명: {project_name}")
    doc.add_paragraph(f"과제번호: {project_no}")
    doc.add_paragraph(f"기간: {period}")
    doc.add_paragraph(f"예산: {budget} 천원\n")

    prompts = [build_prompt(s, project_name, depart_name, project_no, period, budget) for s in sections]
    generated = generate_text_batch(prompts)

    all_refs_flat = []
    for sec, gen_text in zip(sections, generated):
        doc.add_heading(sec['section'], level=1)
        cleaned = clean_generated_text(gen_text)
        refs = search_reference(sec["query"], topk=TOPK)
        validation = validate_output(sec, cleaned, refs)
        doc.add_paragraph(cleaned)
        p = doc.add_paragraph()
        p.add_run("[Eval_Result]").bold = True
        for v in validation:
            doc.add_paragraph(v)
        all_refs_flat.extend(refs)

    add_reference_section(doc, all_refs_flat)
    outpath = "RND_Report.docx"
    doc.save(outpath)
    print(f"[DONE] → {outpath}")

# ==============================
# 12. 실행 예시
# ==============================
if __name__ == "__main__":
    project_name = "VUNO Med-Chest X-ray"
    depart_name  = "산업기술R&D연구기획사업"
    project_no   = "123456789"
    period       = "2023.6.1 ~"
    budget       = "500,000"
    print(f"[INFO] 시작: {project_name} ({depart_name})")
    render_doc(project_name, depart_name, project_no, period, budget)